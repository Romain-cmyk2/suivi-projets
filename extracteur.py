"""
Extracteur de contenu de documents pour pré-remplir le wizard de création de projet.
Supporte : .docx, .xlsx, .pdf, .txt, .md, .csv
Extrait le texte brut puis parse les sections structurées.
"""

import os
import re
from datetime import datetime, date, timedelta


def extraire_texte(filepath):
    """Extrait le texte brut d'un fichier selon son extension."""
    ext = os.path.splitext(filepath)[1].lower()
    texte = ""

    if ext == ".txt" or ext == ".md":
        with open(filepath, "r", encoding="utf-8", errors="replace") as f:
            texte = f.read()

    elif ext == ".csv":
        with open(filepath, "r", encoding="utf-8", errors="replace") as f:
            texte = f.read()

    elif ext == ".docx":
        try:
            from docx import Document
            doc = Document(filepath)
            parties = []
            for para in doc.paragraphs:
                parties.append(para.text)
            # Aussi lire les tableaux
            for table in doc.tables:
                for row in table.rows:
                    ligne = " | ".join(cell.text.strip() for cell in row.cells)
                    parties.append(ligne)
            texte = "\n".join(parties)
        except ImportError:
            raise ImportError("python-docx est requis : pip install python-docx")

    elif ext == ".xlsx" or ext == ".xls":
        try:
            import pandas as pd
            xls = pd.ExcelFile(filepath)
            parties = []
            for sheet_name in xls.sheet_names:
                df = pd.read_excel(xls, sheet_name=sheet_name)
                parties.append(f"=== Feuille: {sheet_name} ===")
                parties.append(df.to_string(index=False))
            texte = "\n".join(parties)
        except ImportError:
            raise ImportError("openpyxl est requis : pip install openpyxl")

    elif ext == ".pdf":
        try:
            from pypdf import PdfReader
            reader = PdfReader(filepath)
            parties = []
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    parties.append(t)
            texte = "\n".join(parties)
        except ImportError:
            try:
                import PyPDF2
                reader = PyPDF2.PdfReader(filepath)
                parties = []
                for page in reader.pages:
                    t = page.extract_text()
                    if t:
                        parties.append(t)
                texte = "\n".join(parties)
            except ImportError:
                raise ImportError("pypdf ou PyPDF2 est requis : pip install pypdf")

    return texte.strip()


def parser_sections(texte):
    """
    Parse le texte pour identifier les sections structurées.
    Cherche des patterns courants : titres, listes, tableaux...
    Retourne un dict avec les champs pré-remplis.
    """
    resultat = {
        "nom": "",
        "description": "",
        "categorie": "Autre",
        "objectifs": [],
        "livrables": "",
        "contraintes": "",
        "taches": [],
    }

    lignes = texte.split("\n")
    lignes = [l.strip() for l in lignes if l.strip()]

    # --- Détection du nom du projet ---
    # Première ligne significative ou ligne après "Projet :" / "Titre :"
    for l in lignes[:10]:
        m = re.match(r"(?:projet|titre|nom(?:\s+du\s+projet)?)\s*[:;]\s*(.+)", l, re.IGNORECASE)
        if m:
            resultat["nom"] = m.group(1).strip()
            break
    if not resultat["nom"] and lignes:
        # Prendre la première ligne courte comme titre
        for l in lignes[:5]:
            cleaned = re.sub(r"^[#=\-*]+\s*", "", l).strip()
            if 3 < len(cleaned) < 120:
                resultat["nom"] = cleaned
                break

    # --- Extraction par sections ---
    section_courante = ""
    contenu_section = {}
    section_patterns = {
        "description": r"(?:^|\b)(?:description|contexte|resume|presentation|introduction|objet)(?:\b|$|\s*:)",
        "objectifs": r"(?:^|\b)(?:objectifs?|buts?|finalite|goals?)(?:\b|$|\s*:)",
        "livrables": r"(?:^|\b)(?:livrables?|deliverables?|resultats?\s+attendus?|outputs?|produits?\s+finis?)(?:\b|$|\s*:)",
        "contraintes": r"(?:^|\b)(?:contraintes?|risques?|limites?|restrictions?|prerequis|dependances?)(?:\b|$|\s*:)",
        "taches": r"(?:^|\b)(?:t[aâ]ches?|actions?\s+[àa]\s+mener|[eé]tapes?|phases?|activit[eé]s?|todo|planning|plan\s+d.action|work\s*packages?|lots?\s+de\s+travaux|macro.?planning)(?:\b|$|\s*:)",
    }

    # Un titre de section doit être court et ressembler à un en-tête, pas à du contenu
    def est_titre_section(ligne, pattern):
        l_clean = re.sub(r"^[\d.)\-#*=]+\s*", "", ligne).strip()
        if len(l_clean) > 60:
            return False
        # Le mot-clé doit représenter l'essentiel de la ligne (pas noyé dans une phrase)
        match = re.search(pattern, l_clean, re.IGNORECASE)
        if not match:
            return False
        # Si le match représente au moins 30% de la ligne, c'est un titre
        ratio = len(match.group()) / len(l_clean)
        return ratio > 0.3

    for l in lignes:
        # Détecter un titre de section
        detected = False
        for section_name, pattern in section_patterns.items():
            if est_titre_section(l, pattern):
                section_courante = section_name
                if section_courante not in contenu_section:
                    contenu_section[section_courante] = []
                detected = True
                break

        if not detected and section_courante:
            contenu_section.setdefault(section_courante, []).append(l)

    # --- Remplir les champs ---

    # Description
    if "description" in contenu_section:
        resultat["description"] = "\n".join(contenu_section["description"])
    elif not resultat["description"]:
        # Prendre les premières lignes comme description
        desc_lignes = []
        for l in lignes[1:15]:
            if len(l) > 10:
                desc_lignes.append(l)
        resultat["description"] = "\n".join(desc_lignes[:5])

    # Objectifs
    if "objectifs" in contenu_section:
        for l in contenu_section["objectifs"]:
            obj = re.sub(r"^[\-\*•●◦▪\d.)\]]+\s*", "", l).strip()
            if obj and len(obj) > 5:
                resultat["objectifs"].append(obj)

    # Livrables
    if "livrables" in contenu_section:
        items = []
        for l in contenu_section["livrables"]:
            item = re.sub(r"^[\-\*•●◦▪\d.)\]]+\s*", "", l).strip()
            if item and len(item) > 3:
                items.append(item)
        resultat["livrables"] = "\n".join(items)

    # Contraintes
    if "contraintes" in contenu_section:
        items = []
        for l in contenu_section["contraintes"]:
            item = re.sub(r"^[\-\*•●◦▪\d.)\]]+\s*", "", l).strip()
            if item and len(item) > 3:
                items.append(item)
        resultat["contraintes"] = "\n".join(items)

    # Tâches
    if "taches" in contenu_section:
        for l in contenu_section["taches"]:
            tache_txt = re.sub(r"^[\-\*•●◦▪\d.)\]]+\s*", "", l).strip()
            if tache_txt and len(tache_txt) > 3:
                # Essayer d'extraire une durée
                duree = 5
                m_duree = re.search(r"(\d+)\s*(?:jour|j\b|day)", tache_txt, re.IGNORECASE)
                if m_duree:
                    duree = int(m_duree.group(1))

                # Essayer d'extraire une priorité
                priorite = "Moyenne"
                if re.search(r"(?:critique|urgent|critical)", tache_txt, re.IGNORECASE):
                    priorite = "Critique"
                elif re.search(r"(?:haute|high|important)", tache_txt, re.IGNORECASE):
                    priorite = "Haute"
                elif re.search(r"(?:basse|low|optionnel|facultat)", tache_txt, re.IGNORECASE):
                    priorite = "Basse"

                # Nettoyer le titre (enlever les annotations de durée/priorité)
                titre = re.sub(r"\s*[\(\[].+?[\)\]]", "", tache_txt).strip()
                titre = re.sub(r"\s*\|\s*.+$", "", titre).strip()
                if titre:
                    resultat["taches"].append({
                        "titre": titre,
                        "priorite": priorite,
                        "assigne": "",
                        "duree_jours": min(duree, 60),
                    })

    # --- Détection de la catégorie ---
    texte_lower = texte.lower()
    if any(w in texte_lower for w in ["developpement", "dev", "logiciel", "application", "code", "api"]):
        resultat["categorie"] = "Developpement"
    elif any(w in texte_lower for w in ["infrastructure", "serveur", "reseau", "deploiement", "migration"]):
        resultat["categorie"] = "Infrastructure"
    elif any(w in texte_lower for w in ["formation", "training", "apprentissage", "montee en competence"]):
        resultat["categorie"] = "Formation"
    elif any(w in texte_lower for w in ["communication", "marketing", "campagne", "evenement"]):
        resultat["categorie"] = "Communication"
    elif any(w in texte_lower for w in ["organisation", "process", "procedure", "reorganisation"]):
        resultat["categorie"] = "Organisation"

    return resultat


def extraire_depuis_excel_structure(filepath):
    """
    Extraction spéciale pour les fichiers Excel structurés
    avec colonnes : Tâche, Priorité, Assigné, Durée, etc.
    """
    import pandas as pd

    try:
        df = pd.read_excel(filepath)
    except Exception:
        return None

    cols_lower = {c: c.lower().strip() for c in df.columns}
    col_map = {}

    # Mapper les colonnes
    for col, col_l in cols_lower.items():
        if any(w in col_l for w in ["tache", "tâche", "titre", "task", "action", "activite", "activité"]):
            col_map["titre"] = col
        elif any(w in col_l for w in ["priorite", "priorité", "priority", "prio"]):
            col_map["priorite"] = col
        elif any(w in col_l for w in ["assigne", "assigné", "responsable", "owner", "qui"]):
            col_map["assigne"] = col
        elif any(w in col_l for w in ["duree", "durée", "jours", "days", "duration"]):
            col_map["duree"] = col
        elif any(w in col_l for w in ["description", "detail", "détail", "commentaire"]):
            col_map["description"] = col

    if "titre" not in col_map:
        return None

    taches = []
    for _, row in df.iterrows():
        titre = str(row.get(col_map["titre"], "")).strip()
        if not titre or titre == "nan":
            continue

        priorite = "Moyenne"
        if "priorite" in col_map:
            p = str(row.get(col_map["priorite"], "")).strip().lower()
            if p in ["critique", "critical", "urgent"]:
                priorite = "Critique"
            elif p in ["haute", "high", "important", "élevée", "elevee"]:
                priorite = "Haute"
            elif p in ["basse", "low", "faible"]:
                priorite = "Basse"

        assigne = ""
        if "assigne" in col_map:
            a = str(row.get(col_map["assigne"], "")).strip()
            if a and a != "nan":
                assigne = a

        duree = 5
        if "duree" in col_map:
            try:
                d = int(float(row.get(col_map["duree"], 5)))
                duree = max(1, min(d, 60))
            except (ValueError, TypeError):
                pass

        taches.append({
            "titre": titre,
            "priorite": priorite,
            "assigne": assigne,
            "duree_jours": duree,
        })

    return taches if taches else None


def extraire_document(filepath):
    """
    Point d'entrée principal.
    Retourne un dict avec tous les champs pré-remplis + le texte brut.
    """
    ext = os.path.splitext(filepath)[1].lower()

    # D'abord essayer l'extraction structurée pour Excel
    if ext in (".xlsx", ".xls"):
        taches_excel = extraire_depuis_excel_structure(filepath)
        if taches_excel:
            texte = extraire_texte(filepath)
            resultat = parser_sections(texte)
            resultat["taches"] = taches_excel
            resultat["texte_brut"] = texte
            resultat["source_fichier"] = os.path.basename(filepath)
            return resultat

    # Extraction texte classique
    texte = extraire_texte(filepath)
    if not texte:
        return None

    resultat = parser_sections(texte)
    resultat["texte_brut"] = texte
    resultat["source_fichier"] = os.path.basename(filepath)
    return resultat
